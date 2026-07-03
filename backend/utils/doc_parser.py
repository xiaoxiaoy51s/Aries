"""文档文本抽取工具：支持 PDF / Word / Excel 文件，抽取纯文本供 AI 通读。

设计目标：无向量化模型时，"暴力阅读全部内容"。
所有解析器统一返回纯文本字符串，由调用方负责截断/分页。
"""
from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# 支持的文档扩展名（小写，含点）
SUPPORTED_DOC_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".doc", ".xls", ".xlsm"}


def is_supported_document(path: Path | str) -> bool:
    """判断文件是否为可解析的文档类型。"""
    ext = Path(path).suffix.lower()
    return ext in SUPPORTED_DOC_EXTENSIONS


def extract_text(path: Path | str) -> str:
    """根据扩展名分发到对应解析器，返回抽取出的纯文本。

    失败时抛出 ValueError（含可读的失败原因），由调用方转成错误响应。
    """
    p = Path(path)
    ext = p.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(p)
    if ext == ".docx":
        return _extract_docx(p)
    if ext in (".xlsx", ".xlsm"):
        return _extract_xlsx(p)
    if ext in (".doc", ".xls"):
        # 旧版二进制格式，python 库无法直接解析，需 LibreOffice 预转换
        raise ValueError(
            f"不支持旧版 .{ext.lstrip('.')} 格式，请先用 LibreOffice 转换为 .{'docx' if ext == '.doc' else 'xlsx'} 后重试"
        )
    raise ValueError(f"不支持的文档类型: {ext}")


def _extract_pdf(path: Path) -> str:
    """使用 pypdf 抽取 PDF 全部文本，按页输出。"""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ValueError("缺少依赖 pypdf，请先安装: pip install pypdf")

    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        for i, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text() or ""
            except Exception as exc:
                text = ""
                logger.warning(f"PDF 第 {i} 页文本抽取失败: {exc}")
            parts.append(f"--- 第 {i} 页 ---")
            parts.append(text)
        result = "\n".join(parts).strip()
        if not result:
            raise ValueError("PDF 未抽取到任何文本（可能是扫描件/纯图片，需要 OCR）")
        return result
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"PDF 解析失败: {exc}")


def _extract_docx(path: Path) -> str:
    """使用 python-docx 抽取 Word 文档段落与表格文本。"""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("缺少依赖 python-docx，请先安装: pip install python-docx")

    try:
        doc = Document(str(path))
        parts: list[str] = []

        # 段落（含标题）
        for para in doc.paragraphs:
            text = para.text
            if text and text.strip():
                parts.append(text)

        # 表格：按行拼接单元格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))

        result = "\n".join(parts).strip()
        if not result:
            raise ValueError("Word 文档未抽取到任何文本（可能是空文档或纯图片）")
        return result
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Word 文档解析失败: {exc}")


def _extract_xlsx(path: Path) -> str:
    """使用 openpyxl 抽取 Excel 各 sheet 的单元格文本。"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ValueError("缺少依赖 openpyxl，请先安装: pip install openpyxl")

    try:
        # read_only 内存友好；data_only 读取公式计算后的值
        wb = load_workbook(str(path), read_only=True, data_only=True)
        parts: list[str] = []
        try:
            for ws in wb.worksheets:
                parts.append(f"=== Sheet: {ws.title} ===")
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if c is None else str(c) for c in row]
                    # 跳过全空行，减少噪音
                    if any(cell.strip() for cell in cells):
                        parts.append("\t".join(cells))
                        row_count += 1
                parts.append(f"（共 {row_count} 行数据）")
                parts.append("")
        finally:
            wb.close()
        result = "\n".join(parts).strip()
        if not result:
            raise ValueError("Excel 文档未抽取到任何数据（可能是空表）")
        return result
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Excel 文档解析失败: {exc}")
